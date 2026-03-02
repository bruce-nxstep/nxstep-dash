"""
cv_parser.py
Extracts raw text from an uploaded CV file (PDF or DOCX).
"""
import os
import pdfplumber
from docx import Document


def extract_text_from_file(filepath: str) -> str:
    """
    Extract plain text from a PDF or DOCX file.
    Returns the full text content as a string.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _extract_from_pdf(filepath)
    elif ext == ".docx":
        return _extract_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def _extract_from_pdf(filepath: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    pages_text = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())
    return "\n\n".join(pages_text)


def _extract_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)
