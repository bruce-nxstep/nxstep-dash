"""
docx_generator.py
Generates an ATS-compliant DOCX CV from structured data,
following the NXSTEP template structure using python-docx.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# ── Colour palette ─────────────────────────────────────────
RGB_BLACK  = RGBColor(0x0D, 0x0D, 0x0D)
RGB_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
RGB_ACCENT = RGBColor(0x2E, 0x40, 0x57)
RGB_MID    = RGBColor(0x4A, 0x55, 0x68)
RGB_LIGHT  = RGBColor(0x71, 0x80, 0x96)
RGB_LINE   = RGBColor(0xCB, 0xD5, 0xE0)


def _set_run_font(run, size_pt: float, bold=False, italic=False, color: RGBColor = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_horizontal_rule(doc: Document, color_hex="CBD5E0", thickness_pt=4):
    """Insert a bottom border on the last paragraph as a visual separator."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    # Add bottom border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness_pt))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_section_header(doc: Document, title: str):
    """Add an uppercase section header with accent color and divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title.upper())
    _set_run_font(run, 10, bold=True, color=RGB_DARK)

    # Add bottom border on section title paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E4057')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_bullet(doc: Document, text: str):
    """Add a bullet point paragraph (ATS-safe: no list style, just bullet char)."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(-0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(f"• {text.strip()}")
    _set_run_font(run, 9.5, color=RGB_BLACK)
    return para


def _add_body(doc: Document, text: str, bold_prefix: str = None, size=9.5):
    """Add a body paragraph, optionally with a bolded prefix."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run_b = para.add_run(bold_prefix)
        _set_run_font(run_b, size, bold=True, color=RGB_BLACK)
        run = para.add_run(f" {text}")
        _set_run_font(run, size, color=RGB_BLACK)
    else:
        run = para.add_run(text)
        _set_run_font(run, size, color=RGB_BLACK)
    return para


def generate_cv_docx(data: dict) -> bytes:
    """
    Generate an ATS-compliant DOCX from structured CV data dict.
    Returns the DOCX as bytes.
    """
    doc = Document()

    # ── Page margins (narrow for more content space) ─────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Defaults ─────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    nom    = data.get("nom_complet") or "Candidat"
    titre  = data.get("titre_poste") or ""
    contact = data.get("contact") or {}

    # ══════════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════════
    # Name
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after  = Pt(2)
    run = p_name.add_run(nom)
    _set_run_font(run, 22, bold=True, color=RGB_BLACK)

    # Job title
    if titre:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(3)
        run = p_title.add_run(titre)
        _set_run_font(run, 12, color=RGB_ACCENT)

    # Contact line
    parts = []
    for key in ("email", "telephone", "localisation", "linkedin", "github"):
        val = contact.get(key)
        if val:
            parts.append(val)
    if parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(4)
        run = p_contact.add_run("  |  ".join(parts))
        _set_run_font(run, 8.5, color=RGB_MID)

    # Heavy divider under header
    _add_horizontal_rule(doc, color_hex="0D0D0D", thickness_pt=8)

    # ══════════════════════════════════════════════════════════
    # PROFIL
    # ══════════════════════════════════════════════════════════
    profil = data.get("profil")
    if profil:
        _add_section_header(doc, "Profil Professionnel")
        _add_body(doc, profil)

    # ══════════════════════════════════════════════════════════
    # COMPÉTENCES
    # ══════════════════════════════════════════════════════════
    competences = data.get("competences") or []
    if competences:
        _add_section_header(doc, "Compétences & Savoir-Faire")
        for cat in competences:
            cat_name = cat.get("categorie") or ""
            items    = cat.get("items") or []
            if cat_name:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after  = Pt(1)
                run = p.add_run(cat_name)
                _set_run_font(run, 9.5, bold=True, color=RGB_ACCENT)
            for item in items:
                if item and item.strip():
                    _add_bullet(doc, item)

    # ══════════════════════════════════════════════════════════
    # OUTILS TECHNIQUES
    # ══════════════════════════════════════════════════════════
    outils = data.get("outils_techniques") or []
    if outils:
        _add_section_header(doc, "Outils & Technologies")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(", ".join(o for o in outils if o))
        _set_run_font(run, 9, color=RGB_MID)

    # ══════════════════════════════════════════════════════════
    # EXPÉRIENCES
    # ══════════════════════════════════════════════════════════
    experiences = data.get("experiences") or []
    if experiences:
        _add_section_header(doc, "Expériences Professionnelles")
        for exp in experiences:
            poste      = exp.get("poste") or ""
            entreprise = exp.get("entreprise") or ""
            periode    = exp.get("periode") or ""
            contexte   = exp.get("contexte") or ""
            mission    = exp.get("mission") or ""
            activites  = exp.get("activites") or []
            env_tech   = exp.get("environnement_technique") or []

            # Role line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            if poste:
                run = p.add_run(poste)
                _set_run_font(run, 10, bold=True, color=RGB_ACCENT)
            if entreprise:
                run = p.add_run(f"  –  {entreprise}")
                _set_run_font(run, 9.5, color=RGB_MID)
            if periode:
                run = p.add_run(f"  ·  {periode}")
                _set_run_font(run, 9, italic=True, color=RGB_LIGHT)

            if contexte:
                _add_body(doc, contexte, bold_prefix="Contexte :")
            if mission:
                _add_body(doc, mission, bold_prefix="Mission :")
            if activites:
                p = doc.add_paragraph()
                run = p.add_run("Activités :")
                _set_run_font(run, 9.5, bold=True, color=RGB_BLACK)
                p.paragraph_format.space_after = Pt(1)
                for act in activites:
                    if act and act.strip():
                        _add_bullet(doc, act)
            if env_tech:
                _add_body(doc, ", ".join(e for e in env_tech if e),
                          bold_prefix="Environnement technique :", size=9)

    # ══════════════════════════════════════════════════════════
    # FORMATION
    # ══════════════════════════════════════════════════════════
    formations = data.get("formations") or []
    if formations:
        _add_section_header(doc, "Formation")
        for f in formations:
            parts = []
            if f.get("diplome"):    parts.append(f["diplome"])
            if f.get("etablissement"): parts.append(f["etablissement"])
            if f.get("annee"):      parts.append(f["annee"])
            if parts:
                _add_body(doc, "  –  ".join(parts))

    # ══════════════════════════════════════════════════════════
    # CERTIFICATIONS
    # ══════════════════════════════════════════════════════════
    certifs = data.get("certifications") or []
    if certifs:
        _add_section_header(doc, "Certifications")
        for c in certifs:
            if c and c.strip():
                _add_bullet(doc, c)

    # ══════════════════════════════════════════════════════════
    # LANGUES
    # ══════════════════════════════════════════════════════════
    langues = data.get("langues") or []
    if langues:
        _add_section_header(doc, "Langues")
        for lang in langues:
            l_name = lang.get("langue") or ""
            l_niv  = lang.get("niveau") or ""
            if l_name:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
                run = p.add_run(l_name)
                _set_run_font(run, 9.5, bold=True, color=RGB_BLACK)
                if l_niv:
                    run2 = p.add_run(f" : {l_niv}")
                    _set_run_font(run2, 9.5, color=RGB_BLACK)

    # ══════════════════════════════════════════════════════════
    # CENTRES D'INTÉRÊT
    # ══════════════════════════════════════════════════════════
    centres = data.get("centres_interet") or []
    if centres:
        _add_section_header(doc, "Centres d'Intérêt")
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(c for c in centres if c))
        _set_run_font(run, 9, color=RGB_MID)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
